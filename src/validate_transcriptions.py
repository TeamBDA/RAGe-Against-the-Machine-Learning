# ── Info for error handling and validation check functions──────────────────────────────────────────────────────
import os
import pandas as pd
import datetime
from docx import Document
from functools import partial

# ── Project Path Setup ──────────────────────────────────────────────────────────────────
BASE_DIR           = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TRANSCRIPTIONS_DIR = os.path.join(BASE_DIR, "data/results")
REPORT_DIR         = os.path.join(BASE_DIR, "documents")
REPORT_FILE        = os.path.join(REPORT_DIR, f"csvreport_{datetime.date.today()}.docx")
CSV_FILE           = os.path.join(TRANSCRIPTIONS_DIR, "enriched_transcripts.csv")
# ── END Project Path Setup ───────────────────────────────────────────────────────────────

#Parameters used for the functions below, we can expand this as we add more functions and checks, and it keeps the parameters in one place for easy reference and maintenance. We can also use this to build a config file later if we want to make it more dynamic and not have to change code to update parameters.
PARAMS = {
    "csv_read": {
        "header": 0,
        "skip_blank_lines": True,
        "na_values": ["", " ", "  "]
    },

    "dtype_checks": {
        "cols_int": ['num_words', 'speech_rate_wps', 'speaker_turn_id', 'total_speaking_time_seconds'],
        "col_bool": 'question_flag'
    },

    "timestamp_checks": {
        "timestamp_col": "timestamp"
    }
}

def get_validation_msgs(success, message):
    """Prints a structured validation message."""
    status = "VALIDATION SUCCESS" if success else "VALIDATION FAILED"
    fullmsg= f"[{status}] {message}"
    print(fullmsg)
    return fullmsg

def log_error(table, msg):
    '''Logs an error message to the report table and returns the full message. But first checks that the table has cells to avoid errors when trying to access cells of a row that may not exist.'''
    fullmsg = get_validation_msgs(False, msg)
    row = table.add_row()
    if row.cells:
        row.cells[0].text = str(fullmsg)
    return fullmsg

def log_success(table, msg):
    ''''Logs a success message to the report table and returns the full message. But first checks that the table has cells to avoid errors when trying to access cells of a row that may not exist.'''
    fullmsg = get_validation_msgs(True, msg)
    row = table.add_row()
    if row.cells:
        row.cells[0].text = str(fullmsg)
    return fullmsg

def get_dtype_checks(df, table, cols_int, col_bool):
    """Data type validation logic."""
    # Arrs to collect all errors and a separate array to collect
    # boolean data type errors to keep separated
    errors      = []
    numeric_ok_cols = []

    try:

        print(f"Data types for the columns are currently: \n {df.dtypes}")

        # Using pandas to check if the expected numeric columns are actually numeric,
        # and if not log an error for each column that is not numeric, and if it is
        # numeric add it to a list of valid numeric columns to proceed with validation.
        # This way we can avoid trying to validate non-numeric columns and just log the
        # error for the missing numeric column.'''
        for col in cols_int:
            # Check if column is missing
            if col not in df.columns:
             msg = f"CRITICAL MISSING COLUMN: The required numeric column '{col}' is missing from the CSV file."
             fullmsg = log_error(table, msg)             
             errors.append(fullmsg)
             continue # Skip to the next column

            if not pd.api.types.is_numeric_dtype(df[col]):
                col_num = df.columns.get_loc(col)+1  # Get the column index for the missing numeric column
                msg = f"Missing expected numeric column: {col} at column number {col_num} instead it is {df[col].dtype}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
                
            else:
                # If the column is numeric, add it to the list of valid numeric columns   
                numeric_ok_cols.append(col)

        # Now use if column is numeric to proceed with validation
        if numeric_ok_cols:
            # Integer/float numeric validation in few lines as possible by using pandas to conert to numeric and identify where its not a number or less than or equal to 0, then we can loop through the invalids and print out the row and column of the invalid entry for reporting
            converted = df[cols_int].apply(pd.to_numeric, errors='coerce')
        
            invalid_masked = converted.isna() | (converted <= 0)

            # stack the masked invalid entries to avoid looping thru entire df
            invalid_entries = invalid_masked.stack()
            invalid_entries = invalid_entries[invalid_entries]

            # Loop through the invalid numeric entries and print out the row and
            # column of the invalid entry for reporting
            for (idx, col), _ in invalid_entries.items():
                val = df.at[idx, col]
                msg = f"Invalid numeric value '{val}' at column {col} and row {idx + 2}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)

        # Vectorised boolean validation for the expected boolean column, 
        # but only if the column is present in the file, otherwise log an 
        # error for missing expected boolean column
        if col_bool not in df.columns:
            msg = (
                "CRITICAL MISSING COLUMN: The required boolean column" 
                f" '{col_bool}' is missing."
            )
            errors.append(log_error(table, msg))
        else:
            vectorised = df[col_bool].astype(str).str.strip().str.lower()

            # Flatten the boolean validation to avoid looping through entire df,
            # and just identify the invalid boolean entries, and log an error for
            # each invalid boolean entry found, and append valid boolean values
            # to a list of validated bools
            valid_true = vectorised.isin(["true", "1", "yes"])
            valid_false = vectorised.isin(["false", "0", "no"])
            invalid_bool_mask = ~(valid_true | valid_false)

            df[col_bool] = valid_true

            # now deal with boolean True, false, 1,0, yes, no data type, again skip header row
            for idx in invalid_bool_mask[invalid_bool_mask].index:
                # Any other outcome for the boolean data type should throw an exception,
                # and append all exceptions found
                val = vectorised.at[idx]  # Define val before using it in the error message
                msg = f"Invalid boolean '{val}' at column {col_bool} and row {idx + 2}"         
                errors.append(log_error(table, msg))
                continue

            if errors:
                #log any errors
                log_error(table, "[VALIDATION FAILED] Data type validation failed.")

            else:
                # Return validation correct of all intended data met the data types
                # requirements in file
                print("\n[VALIDATION SUCCESS] Data validation passed.")

                # Log to the Word table
                msg = (
                    "Data type validation passed successfully"
                     " (Numeric and Boolean types are valid)."
                )
                log_success(table, msg)

    # exception to identify any exceptions that occur while running thru csv, and append to all errors
    except Exception as e:
            msg = f"Validation error: {e}"
            fullmsg = log_error(table, msg)             
            errors.append(fullmsg)

def get_csv_checks(df, table):
    """Validates CSV structure, empties, and row count."""
    # Intiailise row count to none initially
    errors      = []
    rowlength = None

    try:
        # apply a lambda (once-off) function to apply mapping to strip extra spaces in csv columns that may exist
        df = df.apply(
            lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x)
        )

        # find indices of empty cells (NaN) using numpy
        empty_cells = pd.isna(df)
        row_indices, col_indices = empty_cells.to_numpy().nonzero()

        # If there are row indices to traverse       
        if len(row_indices) > 0:
            for row, col in zip(row_indices, col_indices):
                msg = f"\n[VALIDATION FAILED] Data validation failed: Empty cell found at CSV line {row + 2}, Column index: {col}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
        else:
            print("There are no empty columns or cells in the CSV file.")
            # Log to the Word table
            msg = (
                "CSV Structure Check: No empty columns"
                 " or cells found in the CSV file."
            )
            log_success(table, msg)

        # used to get row length input - capture no of rows in csv
        rowlength = len(df)
        print(f"There are {rowlength} data rows in the CSV file.")

        # find if there are more than the expected no of rows
        if rowlength >= 25:
            msg = f"Row count ({rowlength}) is valid (25 or more)."
            fullmsg = log_success(table, msg)             
            
        # capture where there are less than the no of expected rows
        else:
            msg = f"Row count ({rowlength}) is below the minimum required (25)."
            fullmsg = log_error(table, msg)             
            errors.append(fullmsg)

    # catch exceptions where csv is empty
    except pd.errors.EmptyDataError:
        msg = "The CSV file is empty or has a malformed header."
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # or file is not present
    except FileNotFoundError:
        msg =  f"File not found at: {CSV_FILE}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # or any generic exception
    except Exception as e:
        msg = f"An unexpected parsing error occurred: {str(e)}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)

def get_timestamp_checks(df, table, timestamp_col):
    """Timestamp validation logic."""
    errors      = []
    try:
        # intialise index to 0 before we check for timestamps
        idx = 0

        # Check if column is missing
        if timestamp_col not in df.columns:
             msg = (
                 "CRITICAL MISSING COLUMN: The required timestamp column"
                 f" '{timestamp_col}' is missing from the CSV file."
             )
             fullmsg = log_error(table, msg)             
             errors.append(fullmsg)
             return  # Cannot continue without this column

        # Storing a copy of the timestamp so that it is output instead of null NaN val
        df['timestamp_raw'] = df['timestamp']

        # coerce timestamp to meet a readable date/time format
        df['timestamp'] = pd.to_datetime(
            df['timestamp'],
            format="mixed",
            errors="coerce"
        )

        # set variable for a timestamp that does not meet the expected output
        invalid_ts = df[df['timestamp'].isna()]

        # check for timestamp that is not empty first, so test the var above,
        # then identify if the timestamp is in expected raw format
        if not invalid_ts.empty:
            for idx in invalid_ts.index:
                print(f"[INVALID] Row {idx + 2}: timestamp → '{df.loc[idx, 'timestamp_raw']}'")
                msg = f"Timestamp validation failed at row {idx + 2}."
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
            return

        print("All timestamps are valid.")
        # Log to the Word table
        msg = (
            f"Timestamp validation passed successfully." 
            f" Sample formatted entry: {df['timestamp'].iloc[0].strftime('%Y-%m-%d %H:%M:%S')}"
        )
        log_success(table, msg)

        # format timestamp
        dttimeobj = df['timestamp'].iloc[0]
        formatted = dttimeobj.strftime('%Y-%m-%d %H:%M:%S')
        print(f"Timestamp value {formatted}")

    #validate if the file is empty
    except pd.errors.EmptyDataError:
        msg = "The CSV file is empty."
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # capture any generic exceptions
    except Exception as e:
        msg = f"Timestamp validation error: {str(e)}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)

# ── Use for testing validation/error handling functions───────────────────────────────
def run() -> None:
    """
    Pipeline entry point. Loads the saved enriched transcriptions csv, runs all
    validation checks, and saves a dated error report .docx to documents/.
    """
    print(f"Validating CSV file: {CSV_FILE}")
    os.makedirs(REPORT_DIR, exist_ok=True)

    reportdoc = Document()
    reportdoc.add_heading(
        f"Validation Checks Report on {datetime.date.today()}\n"
        f"for file: {CSV_FILE}"
    )

    table = reportdoc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Validation Status & Description"

    validation_pipeline = {
        "csv_checks":   get_csv_checks,
        "dtype_checks": partial(
            get_dtype_checks,
            cols_int=PARAMS["dtype_checks"]["cols_int"],
            col_bool=PARAMS["dtype_checks"]["col_bool"]
        ),
        "timestamp_checks": partial(
            get_timestamp_checks,
            timestamp_col=PARAMS["timestamp_checks"]["timestamp_col"]
        ),
    }

    try:
        READ_PARAMS = PARAMS["csv_read"]
        df_validate = pd.read_csv(CSV_FILE, **READ_PARAMS)

        for func in validation_pipeline.values():
            func(df_validate, table)

    except TypeError as e:
        print(f"[FAILED] Invalid parameter configuration error: {str(e)}")
    except Exception as e:
        print(f"[VALIDATION FAILED] Error: {str(e)}")

    reportdoc.save(REPORT_FILE)
    print(f"✅ Error report saved to {REPORT_FILE}.")


if __name__ == "__main__":
    run()