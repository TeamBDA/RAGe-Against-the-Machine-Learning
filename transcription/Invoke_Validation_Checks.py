# ── Info for error handling and validation check functions──────────────────────────────────────────────────────
import os
import pandas as pd
import datetime
from docx import Document
from functools import partial

# ── Project Path Setup ──────────────────────────────────────────────────────────────────
BASE_DIR           = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TRANSCRIPTIONS_DIR = os.path.join(BASE_DIR, "data")
REPORT_DIR         = os.path.join(BASE_DIR, "documents")
REPORT_FILE        = os.path.join(REPORT_DIR, f"csvreport_{datetime.date.today()}.docx")
CSV_FILE           = os.path.join(TRANSCRIPTIONS_DIR, "corrected_transcripts.csv")
# ── END Project Path Setup ───────────────────────────────────────────────────────────────

PARAMS = {
    "csv_read": {
        "header": 0,
        "skip_blank_lines": True,
        "na_values": ["", " ", "  "]
    },

    "dtype_checks": {
        "cols_int": ['num_words', 'speech_rate_wps', 'speaker_turn_id', 'total_speaking_time_seconds'],
        "col_bool": 'question_flag'
    },

    "timestamp_checks": {
        "timestamp_col": "timestamp"
    }
}

def get_validation_msgs(success, message):
    """Prints a structured validation message."""
    status = "VALIDATION SUCCESS" if success else "VALIDATION FAILED"
    fullmsg= f"[{status}] {message}"
    print(fullmsg)
    return fullmsg

def log_error(table, msg):
    fullmsg = get_validation_msgs(False, msg)
    row = table.add_row()
    if row.cells:
        row.cells[0].text = str(fullmsg)
    return fullmsg

def log_success(table, msg):
    fullmsg = get_validation_msgs(True, msg)
    row = table.add_row()
    if row.cells:
        row.cells[0].text = str(fullmsg)
    return fullmsg

def get_dtype_checks(df, table, cols_int, col_bool):
    """Data type validation logic."""
    # Arrs to collect all errors and a separate array to collect boolean data type errors to keep separated
    errors      = []
    numeric_ok_cols = []

    try:

        print(f"Data types for the columns are currently: \n {df.dtypes}")

        for col in cols_int:
            # Check if column is missing
            if col not in df.columns:
                msg = f"CRITICAL MISSING COLUMN: The required numeric column '{col}' is missing from the CSV file."
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
                continue # Skip to the next column

            if not pd.api.types.is_numeric_dtype(df[col]):
                col_num = df.columns.get_loc(col)+1  # Get the column index for the missing numeric column
                msg = f"Missing expected numeric column: {col} at column number {col_num} instead it is {df[col].dtype}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
            else:
                # If the column is numeric, add it to the list of valid numeric columns   
                numeric_ok_cols.append(col)

        # Now use if column is numeric to proceed with validation
        if numeric_ok_cols:
            # Integer/float numeric validation in few lines as possible by using pandas to conert to numeric and identify where its not a number or less than or equal to 0, then we can loop through the invalids and print out the row and column of the invalid entry for reporting
            converted = df[cols_int].apply(pd.to_numeric, errors='coerce')
        
            invalid_masked = converted.isna() | (converted <= 0)

            # stack the masked invalid entries to avoid looping thru entire df
            invalid_entries = invalid_masked.stack()
            invalid_entries = invalid_entries[invalid_entries]


            # Loop through the invalid numeric entries and print out the row and column of the invalid entry for reporting
            for (idx, col), _ in invalid_entries.items():
                val = df.at[idx, col]
                msg = f"Invalid numeric value '{val}' at column {col} and row {idx + 2}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)

        # Check if the boolean column exists before trying to manipulate it
        if col_bool not in df.columns:
            msg = f"CRITICAL MISSING COLUMN: The required boolean column '{col_bool}' is missing from the CSV file structure."
            errors.append(log_error(table, msg))
        else:
            vectorised = (
                    df[col_bool]
                    .astype(str)
                    .str.strip()
                    .str.lower()
            )

            valid_true = vectorised.isin(["true", "1", "yes"])
            valid_false = vectorised.isin(["false", "0", "no"])
            invalid_bool_mask = ~(valid_true | valid_false)

            df[col_bool] = valid_true

            # now deal with boolean True, false, 1,0, yes, no data type, again skip header row
            for idx in invalid_bool_mask[invalid_bool_mask].index:
                    # Any other outcome for the boolean data type should throw an exception, and append all exceptions found
                    val = vectorised.at[idx]  # Define val before using it in the error message
                    msg = f"Invalid boolean '{val}' at column {col_bool}"         
                    errors.append(log_error(table, msg))
                    continue

        if errors:
            #log any errors
            log_error(table, "Data type validation failed.")

        else:
            # Return validation correct of all intended data met the data types requirements in file
            print("\n[VALIDATION SUCCESS] Data validation passed.")

    # exception to identify any exceptions that occur while running thru csv, and append to all errors
    except Exception as e:
        msg = f"Validation error: {e}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)


def get_csv_checks(df, table):
    """Validates CSV structure, empties, and row count."""
    # Intiailise row count to none initially
    errors      = []
    rowlength = None

    try:
        # apply a lambda (once-off) function to apply mapping to strip extra spaces in csv columns that may exist
        df = df.apply(
            lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x)
        )

        # find indices of empty cells (NaN) using numpy
        empty_cells = pd.isna(df)
        row_indices, col_indices = empty_cells.to_numpy().nonzero()

        # If there are row indices to traverse
        if len(row_indices) > 0:
            for row, col in zip(row_indices, col_indices):
                msg = f"\n[VALIDATION FAILED] Data validation failed: Empty cell found at CSV line {row + 2}, Column index: {col}"
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
        else:
            print("There are no empty columns or cells in the CSV file.")

        # used to get row length input - capture no of rows in csv
        rowlength = len(df)
        print(f"There are {rowlength} data rows in the CSV file.")

        # find if there are more than the expected no of rows
        if rowlength >= 25:
            msg = f"Row count ({rowlength}) is valid (25 or more)."
            fullmsg = log_success(table, msg)             
            
        # capture where there are less than the no of expected rows
        else:
            msg = f"Row count ({rowlength}) is below the minimum required (25)."
            fullmsg = log_error(table, msg)             
            errors.append(fullmsg)

    # catch exceptions where csv is empty
    except pd.errors.EmptyDataError:
        msg = "The CSV file is empty or has a malformed header."
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # or file is not present
    except FileNotFoundError:
        msg =  f"File not found at: {CSV_FILE}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # or any generic exception
    except Exception as e:
        msg = f"An unexpected parsing error occurred: {str(e)}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)


def get_timestamp_checks(df, table, timestamp_col):
    """Timestamp validation logic."""
    errors      = []
    try:
        # intialise index to 0 before we check for timestamps
        idx = 0

        # check for occurence of the timestamp column in file
        if 'timestamp' not in df.columns:
            msg =  f"A timestamp column could not be found."
            fullmsg = log_error(table, msg)             
            errors.append(fullmsg)
            return

        # Storing a copy of the timestamp so that it is output instead of null NaN val
        df['timestamp_raw'] = df['timestamp']

        # coerce timestamp to meet a readable date/time format
        df['timestamp'] = pd.to_datetime(
            df['timestamp'],
            format="mixed",
            errors="coerce"
        )

        # set variable for a timestamp that does not meet the expected output
        invalid_ts = df[df['timestamp'].isna()]

        # check for timestamp that is not empty first, so test the var above, then identify if the timestamp is in expected raw format
        if not invalid_ts.empty:
            for idx in invalid_ts.index:
                print(f"[INVALID] Row {idx + 2}: timestamp → '{df.loc[idx, 'timestamp_raw']}'")
                msg =  "Timestamp validation failed."
                fullmsg = log_error(table, msg)             
                errors.append(fullmsg)
            return

        print("All timestamps are valid.")
        print(df['timestamp'])

        # format timestamp
        dttimeobj = df['timestamp'].iloc[0]
        formatted = dttimeobj.strftime('%Y-%m-%d %H:%M:%S')
        print(f"Timestamp value {formatted}")

    #validate if the file is empty
    except pd.errors.EmptyDataError:
        msg = "The CSV file is empty."
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)
    # capture any generic exceptions
    except Exception as e:
        msg = f"Timestamp validation error: {str(e)}"
        fullmsg = log_error(table, msg)             
        errors.append(fullmsg)

# ── Use for testing validation/error handling functions───────────────────────────────
if __name__ == "__main__":

    print(f"Validating CSV file: {CSV_FILE}")

    """Generate report from CSV file."""
    # Create a docx (Word) doc for reporting
    reportdoc = Document()
    # Add filename to report
    reportdoc.add_heading(f"Validation Checks Report on {datetime.date.today()} \n for file @: \n  {CSV_FILE}")

    # Create a table with 1 header and all rows and columns for the error report generated
    table = reportdoc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Error Description"

    """Run pipeline and generate report we will move this to the main pipeline script for the project later"""
    validation_pipeline = {
        "dtype_checks": partial(get_dtype_checks,
                                cols_int=PARAMS["dtype_checks"]["cols_int"], 
                                col_bool=PARAMS["dtype_checks"]["col_bool"]),
        "csv_checks": get_csv_checks,
        "timestamp_checks": partial(get_timestamp_checks,
                                timestamp_col=PARAMS["timestamp_checks"]["timestamp_col"])
    }

    """Run read file exactly once within main by referencing different param set per function"""
    try:
      # Build unified read parameters
      READ_PARAMS = PARAMS["csv_read"]
      df = pd.read_csv(CSV_FILE, **READ_PARAMS)  # Read all data as string to handle type validation in functions

      '''Run all functions on df'''
      for name, func in validation_pipeline.items():
        func(df, table)

    except TypeError as e:
           print (f"[{PARAMS}] [FAILED]. Invalid parameter configuration error: {str(e)}")
    except Exception as e:
           print (f"[{PARAMS}] [VALIDATION FAILED]. Error: {str(e)}")

    # Save the final report
    reportdoc.save(REPORT_FILE)
